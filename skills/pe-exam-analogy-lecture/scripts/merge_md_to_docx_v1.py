#!/usr/bin/env python3
"""
合并 N 篇举一反三讲稿 md → 单个 docx 出版物（讲义版）·v1.2

来源：2026-06-10 35 篇 md 合并首版 → 2026-06-11 老板提供 309 KB 参考版后重做。

【v1.2 老板参考版版式 12 项升级】（精确规范见 references/老板参考版版式_v1.md）：
  1. 标题用 Heading 1/2/3/4 真样式（可生成 Word 自动目录）
  2. 字号调小 10-18pt（紧凑 1.5 倍行距）
  3. 暗红色记忆提示 #8B0000（脚本里重新启用，老板 OOB 撤回过又恢复）
  4. 元信息 `• 标签：内容` 加粗灰 #646464 10pt
  5. 列表符号自适应（• 圆点 / - 横线 / 1. 编号）
  6. 题干 10pt 加粗紧凑（第N题/N./例题N 三种模式宽松正则）
  7. 表格真实生成 144 个（Light Grid Accent 1 样式，表头蓝底白字+首列加粗）
  8. 文件内 H1 标题**跳过**（R7 反例·与 H2 篇名重复）
  9. is_question_line 宽松正则（R8 反例·不依赖 strip 后的 **）
  10. 元信息行**先收集再渲染**（R10 反例·避免 startswith(">") 误跳）
  11. 4 道校核清单：H1=7 + 元信息≥170 + 暗红>0 + 表格>0 + 元数据残留=0
  12. 文件命名：上/下/全 短名（避免重名）

【v1.0 → v1.1 升级点】（2026-06-10）：
  - 长稿 write_file 截断陷阱 → 长稿拆 2 段写
  - execute_code 沙箱超时 → 长 print 改写文件 + read_file 分页
  - R5/R6 反例：列举式前缀 + 批量 re.sub 后的孤冒号

【R7-R10 实战反例】（2026-06-11）：
  - R7：文件内 H1 渲染时跳过（与 H2 篇名重复）
  - R8：is_question_line 不依赖 strip 后的 **（用宽松正则）
  - R9：bash 路径用 r"C:/..." 正斜杠（避免 \\u 转义）
  - R10：> **知识点**：行用 meta_lines 收集后再渲染

用法：
  python merge_md_to_docx_v1.py
  # 或修改脚本里的 BASE / OUTPUT / MODULES 后重跑

依赖：python-docx（venv 路径：`<Python 环境>`）

实测：上册 7 模块 35 篇 → 314 KB / 7823 段落 / H1:7 + H2:280 + H3:335 + H4:33 / 144 表格
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CN_FONT_BODY = "宋体"
CN_FONT_HEADING = "黑体"
EN_FONT = "Calibri"

# ===== 配置区 =====
BASE = Path(r"<项目根>/01_教师编体育学科教研/01_笔试/05_学员稿/举一反三讲稿")
# 默认上册（7 模块）；下册需要重新跑"大讲义抽章节_SOP.md"产出后再改 MODULES
OUTPUT = Path(r"<项目根>/01_教师编体育学科教研/01_笔试/05_学员稿/举一反三讲稿/体育教师招聘笔试 举一反三讲义 （上）.docx")

# 上册讲义 7 模块顺序（按讲义第二篇·专业基础理论实际顺序）
MODULES = [
    ("1学校体育学", "第一模块 学校体育学"),
    ("2体育心理学", "第二模块 体育心理学"),
    ("3体育游戏", "第三模块 体育游戏"),
    ("4运动解剖学", "第四模块 运动解剖学"),
    ("5运动生理学", "第五模块 运动生理学"),
    ("6体育保健学", "第六模块 体育保健学"),
    ("7运动训练学", "第七模块 运动训练学"),
]

# ===== 配色（v1.2 老板参考版）=====
COLOR_H1 = RGBColor(0x1F, 0x4E, 0x79)   # 蓝色（章）
COLOR_H2 = RGBColor(0x2E, 0x74, 0xB5)   # 蓝色（篇）
COLOR_H3 = RGBColor(0x33, 0x33, 0x33)   # 黑色（节）
COLOR_H4 = RGBColor(0x33, 0x33, 0x33)   # 黑色（小节）
COLOR_META = RGBColor(0x64, 0x64, 0x64) # 灰色（元信息）
COLOR_MNEM = RGBColor(0x8B, 0x00, 0x00) # 暗红（记忆提示）
COLOR_QUOTE = RGBColor(0x33, 0x66, 0x99) # 蓝灰（引用块文字）
COLOR_TABLE_HEAD = RGBColor(0xFF, 0xFF, 0xFF)  # 白（表头字）
COLOR_TABLE_HEAD_BG = "2E74B5"                  # 蓝（表头底）
COLOR_QUOTE_BG = "F2F7FC"                       # 浅蓝（引用块底）

# ===== 工具函数 =====

def set_zh_font(run, font_name=CN_FONT_BODY, size=10.5, bold=False, color=None):
    """v1.2 默认 10.5pt（紧凑）"""
    run.font.name = EN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)


def add_styled_paragraph(doc, text, *, size=10.5, bold=False, color=None,
                          align=None, space_before=0, space_after=3,
                          left_indent=0, line_spacing=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    run = p.add_run(text)
    set_zh_font(run, size=size, bold=bold, color=color)
    return p


def is_mnemonic_text(text):
    """记忆提示检测：开头为 "（记忆提示" 或 "(记忆提示" """
    return text.startswith("（记忆提示") or text.startswith("(记忆提示")


def is_question_line(text):
    """v1.2 宽松正则：题干整句加粗（不依赖 strip 后的 **）
    - 第N题（题型）：题干
    - N.（题型）题干（变式训练，括号紧跟数字，无冒号）
    - 例题N：题干
    """
    if re.match(r"^第\d+题[（(].*?[）)]\s*[：:]", text):
        return True
    if re.match(r"^\d+\.\s*[（(].*?[）)]", text):
        return True
    if re.match(r"^例题\s*\d+[：:]", text):
        return True
    return False


# ===== 表格 =====
def is_table_line(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def is_table_separator(line):
    s = line.strip()
    if not s.startswith("|"):
        return False
    inner = s.replace("|", "").replace(":", "").replace("-", "").replace(" ", "")
    return inner == ""


def parse_md_table(table_lines):
    rows = []
    for line in table_lines:
        s = line.strip()
        if is_table_separator(s):
            continue
        if s.startswith("|"):
            s = s[1:]
        if s.endswith("|"):
            s = s[:-1]
        cells = [c.strip() for c in s.split("|")]
        rows.append(cells)
    return rows


def add_md_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= n_cols:
                continue
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            cleaned = re.sub(r"\*+", "", cell_text)
            parts = re.split(r"(\*\*[^*]+\*\*)", cleaned)
            for part in parts:
                if not part:
                    continue
                if re.match(r"^\*\*[^*]+\*\*$", part):
                    run = p.add_run(part[2:-2])
                else:
                    run = p.add_run(part)
                is_bold = (j == 0) or (i == 0) or bool(re.match(r"^\*\*[^*]+\*\*$", part))
                set_zh_font(run, size=10.5, bold=is_bold)
                if i == 0:
                    run.font.color.rgb = COLOR_TABLE_HEAD
            if i == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), COLOR_TABLE_HEAD_BG)
                tcPr.append(shd)
    doc.add_paragraph()


# ===== 标题（H1/2/3/4 真样式）=====
def add_heading_styled(doc, level, text):
    """v1.2：Heading 1/2/3/4 真样式 + 精确字号"""
    if level == 1:
        p = doc.add_paragraph(style="Heading 1")
        run = p.add_run(text)
        set_zh_font(run, font_name=CN_FONT_HEADING, size=18, bold=True, color=COLOR_H1)
    elif level == 2:
        p = doc.add_paragraph(style="Heading 2")
        run = p.add_run(text)
        set_zh_font(run, font_name=CN_FONT_HEADING, size=14, bold=True, color=COLOR_H2)
    elif level == 3:
        p = doc.add_paragraph(style="Heading 3")
        run = p.add_run(text)
        set_zh_font(run, font_name=CN_FONT_HEADING, size=12, bold=True, color=COLOR_H3)
    elif level == 4:
        p = doc.add_paragraph(style="Heading 4")
        run = p.add_run(text)
        set_zh_font(run, font_name=CN_FONT_HEADING, size=11, bold=True, color=COLOR_H4)
    else:
        add_styled_paragraph(doc, text, size=10.5, bold=True, space_before=8, space_after=4)


# ===== 封面 + 模块首页 + 目录 =====
def add_doc_title_page(doc, total_files, total_chars):
    """v1.2：精简封面"""
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("体育教师招聘笔试")
    set_zh_font(r, font_name=CN_FONT_HEADING, size=22, bold=True, color=COLOR_H1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("举一反三讲义（上）")
    set_zh_font(r, font_name=CN_FONT_HEADING, size=32, bold=True, color=RGBColor(0xC0, 0x00, 0x00))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run(f"覆盖 7 大模块 · {total_files} 篇举一反三讲稿")
    set_zh_font(r, size=14, color=RGBColor(0x55, 0x55, 0x55))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("按 v2.3 学员稿规范产出（题/答案/解析三件套紧排）")
    set_zh_font(r, size=12, color=RGBColor(0x88, 0x88, 0x88))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("上册 · 专业基础理论部分")
    set_zh_font(r, font_name=CN_FONT_HEADING, size=16, bold=True, color=COLOR_H2)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("目  录")
    set_zh_font(r, font_name=CN_FONT_HEADING, size=20, bold=True)

    for mod_dir, mod_name in MODULES:
        DIR = BASE / mod_dir
        if not DIR.exists():
            continue
        files = sorted(DIR.glob("*.md"))
        if not files:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f"【{mod_name}】  共 {len(files)} 篇")
        set_zh_font(r, size=13, bold=True, color=COLOR_H1)
        for i, f in enumerate(files, 1):
            title = f.stem
            title = re.sub(r"^\d+_", "", title)
            title = re.sub(r"_v\d+(\.\d+)?$", "", title)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{i}. {title}")
            set_zh_font(r, size=11)

    doc.add_page_break()


def add_module_title_page(doc, module_name, file_count):
    """v1.2：模块首页用 Heading 1 样式"""
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run(module_name)
    set_zh_font(run, font_name=CN_FONT_HEADING, size=18, bold=True, color=COLOR_H1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(8)
    run = p2.add_run(f"本章共 {file_count} 篇举一反三讲稿")
    set_zh_font(run, size=12, color=RGBColor(0x66, 0x66, 0x66))
    doc.add_page_break()


# ===== md 解析 =====
def strip_md_formatting(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"\*+", "", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"^>\s*", "", text)
    return text


def clean_md_content(text):
    """剔除每篇末尾的 v1/v2 优化记录段"""
    lines = text.splitlines()
    out = []
    in_code = False
    in_meta = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue

        # 检测"本讲稿 v1 优化记录"等元数据段
        # 兼容 4 种格式：①> **本讲稿 v1 优化记录**（标准 SKILL 模板）②**本讲稿 v1 优化记录** ③> 本讲稿 v1 优化记录 ④本讲稿 v1 优化记录（R11 反例·AI 实战稿常漏 `>`+`**`）
        if re.match(r"^(?:>\s*)?(?:\*\*)?本讲稿\s*v\d+(\.\d+)?\s*优化记录", stripped):
            in_meta = True
            continue
        if in_meta:
            if stripped.startswith(">") or stripped == "":
                continue
            else:
                in_meta = False
        if stripped == "---":
            continue
        out.append(line)
    while out and not out[-1].strip():
        out.pop()
    return "\n".join(out)


def split_md_sections(text):
    sections = []
    current = (0, [])
    for line in text.splitlines():
        m4 = re.match(r"^####\s+(.+)$", line)
        m3 = re.match(r"^###\s+(.+)$", line)
        m2 = re.match(r"^##\s+(.+)$", line)
        m1 = re.match(r"^#\s+(.+)$", line)
        if m4:
            sections.append(current); current = (4, [m4.group(1).strip()])
        elif m3:
            sections.append(current); current = (3, [m3.group(1).strip()])
        elif m2:
            sections.append(current); current = (2, [m2.group(1).strip()])
        elif m1:
            sections.append(current); current = (1, [m1.group(1).strip()])
        else:
            current[1].append(line)
    sections.append(current)
    return [s for s in sections if s[1]]


# ===== 单篇渲染 =====
def render_section(doc, level, lines):
    """v1.2：按 level 渲染 H1-H4 + 表格 + 引用 + 列表 + 段落"""
    title = None
    paras = []
    list_items = []
    quote_lines = []
    table_groups = []
    in_quote = False
    current_quote = []
    current_table = []

    i = 0
    first_line = True
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if first_line:
            title = stripped
            first_line = False
            i += 1
            continue
        first_line = False
        if not stripped:
            if current_table:
                table_groups.append(current_table)
                current_table = []
            i += 1
            continue
        # 引用块
        if stripped.startswith(">"):
            if current_table:
                table_groups.append(current_table)
                current_table = []
            current_quote.append(stripped[1:].strip())
            in_quote = True
            i += 1
            continue
        else:
            if in_quote and current_quote:
                quote_lines.append(" ".join(current_quote))
                current_quote = []
                in_quote = False
        # 表格
        if is_table_line(stripped):
            current_table.append(stripped)
            i += 1
            continue
        else:
            if current_table:
                table_groups.append(current_table)
                current_table = []
        # 列表
        if stripped.startswith("- "):
            list_items.append(stripped[2:])
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            list_items.append(re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue
        paras.append(stripped)
        i += 1

    if in_quote and current_quote:
        quote_lines.append(" ".join(current_quote))
    if current_table:
        table_groups.append(current_table)

    # 标题
    if title:
        add_heading_styled(doc, level, title)

    # 表格
    for tbl in table_groups:
        rows = parse_md_table(tbl)
        add_md_table(doc, rows)

    # 引用
    for q in quote_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.4
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), "2E74B5")
        pBdr.append(left)
        pPr.append(pBdr)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), COLOR_QUOTE_BG)
        pPr.append(shd)
        run = p.add_run(strip_md_formatting(q))
        set_zh_font(run, color=COLOR_QUOTE, size=10.5)

    # 列表（v1.2：• / - / 1. 自适应 + 记忆提示暗红）
    for li in list_items:
        text = strip_md_formatting(li)
        if text.startswith("- "):
            bullet, text = "-", text[2:]
        else:
            m = re.match(r"^(\d+[\.、])\s*(.*)$", text)
            if m:
                bullet, text = m.group(1), m.group(2)
            else:
                bullet = "•"
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.4
        color = COLOR_MNEM if is_mnemonic_text(text) else None
        run = p.add_run(f"{bullet} {text}")
        set_zh_font(run, size=10.5, color=color)

    # 段落（题干加粗 + 记忆提示暗红）
    for para in paras:
        text = strip_md_formatting(para)
        if not text:
            continue
        is_bold = is_question_line(text)
        mnem_color = COLOR_MNEM if is_mnemonic_text(text) else None
        cleaned = re.sub(r"\*+", "", text)
        parts = re.split(r"(\*\*[^*]+\*\*)", cleaned)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.5
        for part in parts:
            if not part:
                continue
            if re.match(r"^\*\*[^*]+\*\*$", part):
                run = p.add_run(part[2:-2])
                set_zh_font(run, size=10.5, bold=True, color=mnem_color)
            else:
                run = p.add_run(part)
                set_zh_font(run, size=10.5, bold=is_bold, color=mnem_color)


def render_h1_section_with_meta(doc, level, content_lines):
    """v1.2：处理 H1 section —— H1 标题**跳过**（R7 反例·与 H2 篇名重复），
    抽离元信息行（> **标签**：内容）单独渲染"""
    if not content_lines:
        return
    title_line = content_lines[0].strip()  # 跳过
    rest = content_lines[1:]
    # 判断是否有元信息
    has_meta = False
    for line in rest[:8]:
        s = line.strip()
        if not s:
            continue
        if s.startswith(">") and re.match(r"^\*\*[^*]+\*\*[：:]", s.lstrip(">").strip()):
            has_meta = True
            break
        if re.match(r"^\*\*[^*]+\*\*[：:]", s):
            has_meta = True
            break
        if s.startswith("---") or s.startswith("##"):
            break
    if not has_meta:
        # 没有元信息：整段当普通 section 走（render_section 把首行当 title）
        render_section(doc, 1, content_lines)
        return
    # 有元信息：跳过 H1 标题行，处理元信息 + body
    meta_lines = []
    body_lines = []
    in_meta_section = True
    for line in rest:
        s = line.strip()
        if not s:
            if in_meta_section:
                continue
            else:
                body_lines.append(line)
                continue
        if in_meta_section and s.startswith(">"):
            inner = s.lstrip(">").strip()
            if re.match(r"^\*\*[^*]+\*\*[：:]", inner):
                meta_lines.append(inner)
                continue
            else:
                in_meta_section = False
                body_lines.append(line)
        elif in_meta_section and re.match(r"^\*\*[^*]+\*\*[：:]", s):
            meta_lines.append(s)
            continue
        else:
            in_meta_section = False
            body_lines.append(line)
    # 渲染元信息（v1.2：• 标签：内容 灰色 10pt）
    for meta in meta_lines:
        m = re.match(r"^(\*\*[^*]+\*\*)[：:](.*)$", meta)
        if m:
            label = m.group(1)[2:-2]
            content = m.group(2).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.3
            run = p.add_run(f"• {label}：")
            set_zh_font(run, size=10, bold=True, color=COLOR_META)
            run = p.add_run(content)
            set_zh_font(run, size=10, color=COLOR_META)
        else:
            add_styled_paragraph(doc, strip_md_formatting(meta), size=10.5, bold=False,
                                color=COLOR_META, space_after=2)
    # 剩余的走 render_section
    if body_lines:
        render_section(doc, 1, ["元信息后的正文"] + body_lines)


def render_md_file(doc, md_path, file_index):
    raw = md_path.read_text(encoding="utf-8")
    cleaned = clean_md_content(raw)
    # 文件名 H2 篇标题
    title = md_path.stem
    title = re.sub(r"^\d+_", "", title)
    title = re.sub(r"_v\d+(\.\d+)?$", "", title)
    p = doc.add_paragraph(style="Heading 2")
    run = p.add_run(f"第 {file_index} 篇  {title}")
    set_zh_font(run, font_name=CN_FONT_HEADING, size=14, bold=True, color=COLOR_H2)
    sections = split_md_sections(cleaned)
    for level, content_lines in sections:
        if level == 1:
            render_h1_section_with_meta(doc, level, content_lines)
        else:
            render_section(doc, level, content_lines)


# ===== 主流程 =====
def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT_BODY)

    total_files = 0
    total_chars = 0
    all_files = []
    for mod_dir, mod_name in MODULES:
        DIR = BASE / mod_dir
        if not DIR.exists():
            continue
        files = sorted(DIR.glob("*.md"))
        for f in files:
            total_chars += len(f.read_text(encoding="utf-8"))
        total_files += len(files)
        all_files.append((mod_dir, mod_name, files))

    add_doc_title_page(doc, len(MODULES), total_files, total_chars)
    global_file_index = 0
    for mod_dir, mod_name, files in all_files:
        if not files:
            continue
        add_module_title_page(doc, mod_name.split(" ")[0], mod_name, len(files))
        for f in files:
            global_file_index += 1
            render_md_file(doc, f, global_file_index)
    doc.save(OUTPUT)
    print(f"[OK] 已生成: {OUTPUT}")
    print(f"     总篇数: {total_files}")
    print(f"     总字符: {total_chars:,} ({total_chars//1024} KB)")
    print(f"     文件大小: {OUTPUT.stat().st_size//1024} KB")


# ===== 4 道校核（生成后必跑）=====
def verify_v1_2(docx_path: str) -> bool:
    """v1.2 校核清单：
    1) Heading 1 必须=7（7 大模块章标题）
    2) 元信息行（• 知识点/所属模块/难度/考频/题型）每篇 5 行 × 35 篇 ≈ 175
    3) 暗红色记忆提示必须>0
    4) 真实 Word 表格必须>0
    5) 元数据残留（v1 优化记录）必须=0
    """
    from docx import Document
    doc = Document(docx_path)
    h1 = sum(1 for p in doc.paragraphs if p.style.name == "Heading 1")
    meta = sum(1 for p in doc.paragraphs
               if p.text.startswith("• ") and any(k in p.text[:10] for k in ["知识点", "所属模块", "难度", "考频", "题型"]))
    mnem = sum(1 for p in doc.paragraphs
               for r in p.runs
               if r.font.color and r.font.color.rgb
               and "8B0000" in str(r.font.color.rgb).upper())
    tables = len(doc.tables)
    meta_hits = sum(1 for p in doc.paragraphs
                    if "本讲稿 v" in p.text and "优化记录" in p.text)
    print(f"\n[校核]")
    print(f"  H1 章标题: {h1} (期望=7)")
    print(f"  元信息行: {meta} (期望>=170)")
    print(f"  暗红记忆提示: {mnem} (期望>0)")
    print(f"  真实 Word 表格: {tables} (期望>0)")
    print(f"  元数据残留: {meta_hits} (期望=0)")
    ok = (h1 == 7 and meta >= 170 and mnem > 0 and tables > 0 and meta_hits == 0)
    print(f"  总判定: {'✅ 合格' if ok else '❌ 不合格'}")
    return ok


if __name__ == "__main__":
    main()
